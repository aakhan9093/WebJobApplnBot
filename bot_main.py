import os
import sys
import re
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup

# ─── Add src to path for email helper ───────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent / "src"))
from utils.email_helper import send_job_bot_email

# ─── Load environment ────────────────────────────────────────────────────────
load_dotenv()


# ────────────────────────────────────────────────────────────────────────────
# HELPERS
# ────────────────────────────────────────────────────────────────────────────

def scrape_url(url: str) -> str:
    """Fetch plain text from a URL."""
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n").strip()
    except Exception as e:
        print(f"  ⚠️  Could not scrape URL: {e}")
        return ""


def read_multiline_input(prompt_text: str) -> str:
    """Read multi-line paste from the user (end with a blank line)."""
    print(prompt_text)
    print("  (Paste your text. When done, press Enter on a BLANK line to finish)\n")
    lines = []
    while True:
        try:
            line = input()
            if line == "" and lines and lines[-1] == "":
                break
            lines.append(line)
        except EOFError:
            break
    return "\n".join(lines).strip()


def get_job_description() -> str:
    print("\n" + "=" * 50)
    print("  STEP 1 — JOB DESCRIPTION INPUT")
    print("=" * 50)
    print("  1. Provide a URL (the bot will scrape it)")
    print("  2. Paste JD text directly")
    print()
    choice = input("  Select option (1 or 2): ").strip()

    if choice == "1":
        url = input("  Enter JD URL: ").strip()
        print("  Scraping…")
        text = scrape_url(url)
        if text and len(text) > 200:
            print(f"  ✅ Scraped {len(text)} characters.")
            return text
        else:
            print(f"  ⚠️  Scraping failed or content too short ({len(text) if text else 0} chars).")
            print("     This site might be dynamic (JS-heavy). Please paste it manually below.")
            return read_multiline_input("  Paste the Job Description below:")
    else:
        return read_multiline_input("  Paste the Job Description below:")


def read_pdf(file_path: Path) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except ImportError:
        print("  ⚠️  Library 'pypdf' is missing. Cannot read PDF files.")
        return ""
    except Exception as e:
        print(f"  ⚠️  Error reading PDF: {e}")
        return ""


def get_cv() -> str:
    print("\n" + "=" * 50)
    print("  STEP 2 — CANDIDATE CV")
    print("=" * 50)

    data_dir = Path(__file__).parent / "data"
    cv_files = sorted(
        list(data_dir.glob("*.md")) +
        list(data_dir.glob("*.txt")) +
        list(data_dir.glob("*.pdf"))
    )

    if cv_files:
        print("  Existing CVs found in /data:")
        for i, f in enumerate(cv_files, 1):
            print(f"    {i}. {f.name}")
        print(f"    {len(cv_files) + 1}. Paste a new CV instead")
        print()
        raw = input(f"  Select option (1–{len(cv_files) + 1}): ").strip()
        try:
            idx = int(raw)
        except ValueError:
            idx = len(cv_files) + 1

        if 1 <= idx <= len(cv_files):
            chosen = cv_files[idx - 1]
            print(f"  ✅ Using: {chosen.name}")
            
            # Use appropriate reader based on extension
            if chosen.suffix.lower() == '.pdf':
                return read_pdf(chosen)
            else:
                with open(chosen, "r", encoding="utf-8") as f:
                    return f.read().strip()
    else:
        print("  No CVs found in /data folder.")

    return read_multiline_input("  Paste your CV text below:")


# ────────────────────────────────────────────────────────────────────────────
# DOCX EXPORT — OVERHAULED FOR PREMIUM FORMATTING
# ────────────────────────────────────────────────────────────────────────────

def build_docx(title: str, content: str, filename: str) -> str:
    """Creates a professionally formatted Word doc from AI markdown-like text."""
    doc = Document()
    
    # 🎨 Set Professional Branding Colors
    DEEP_BLUE = RGBColor(0, 51, 102)
    DARK_GREY = RGBColor(64, 64, 64)

    # 📏 Define Styles
    style = doc.styles["Normal"]
    style.font.name = "Outfit" if "Outfit" in [f.name for f in doc.styles] else "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GREY

    # 🏠 Heading 0 - Full Title (Optional)
    if title:
        h0 = doc.add_heading(title.upper(), level=0)
        h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Heading 0 doesn't take RGBColor directly in some docx versions, so we use the first run
        for run in h0.runs:
            run.font.color.rgb = DEEP_BLUE
        doc.add_paragraph()  # spacer

    # Clean the content of the AI-specific markers
    content = content.replace("---FILE_SPLIT---", "").strip()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # 🏷️ Heading Detection (and removal of markers)
        if line.startswith("### "):
            p = doc.add_heading(line[4:].replace("#", "").strip(), level=3)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].replace("#", "").strip(), level=2)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].replace("#", "").strip(), level=1)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE

        # 📍 List Item Detection
        elif line.startswith(("- ", "• ", "* ")) and not line.startswith("**"):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            process_line_with_formatting(p, text)
            
        # 📝 Regular Paragraph with Bold Processing
        else:
            p = doc.add_paragraph()
            process_line_with_formatting(p, line)

    out = Path(__file__).parent / filename
    doc.save(out)
    return str(out)

def process_line_with_formatting(paragraph, text):
    """Handles **bold** markers and other inline formatting while removing the markers."""
    # Regex to catch **bold** text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Remove the markers and add as bold
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
        else:
            # Regular text, but also clean up any stray * markers
            clean_part = part.replace("*", "").strip()
            if clean_part:
                paragraph.add_run(f" {clean_part} " if part != parts[0] else clean_part)


# ────────────────────────────────────────────────────────────────────────────
# MAIN FLOW
# ────────────────────────────────────────────────────────────────────────────

def run():
    jd_text = get_job_description()
    cv_text = get_cv()

    if not jd_text or not cv_text:
        print("\n❌ Both JD and CV are required. Exiting.")
        return

    # ── Load prompt logic ────────────────────────────────────────────────────
    import datetime
    today = datetime.date.today().strftime("%B %d, %Y")

    prompt_file = Path(__file__).parent / "prompts" / "career_transition_logic.md"
    with open(prompt_file, "r", encoding="utf-8") as f:
        prompt_logic = f.read()

    full_prompt = f"""{prompt_logic}

===== CONTEXT =====
TODAY'S DATE: {today}

===== INPUT 1 — JOB DESCRIPTION =====
{jd_text}

===== INPUT 2 — CANDIDATE CV =====
{cv_text}

===== OUTPUT FORMAT INSTRUCTIONS (CRITICAL) =====
You MUST separate your response into FIVE distinct sections using these EXACT labels. Do not nest them. Do not include extra text outside these tags.

[METADATA]
Title: [Extract Job Title]
Company: [Extract Company Name]
Posted: [Look for the date of posting in the JD. If not found, reason from content or use '{today}']
[/METADATA]

[START_MODIFIED_CV]
[STRICT FACTUAL INTEGRITY RULE: Provide ONLY the fully tailored Modified CV here. You MUST NOT invent, make up, or hallucinate any skills, companies, or experiences that are not present in Input 2 (Candidate CV). Your task is to ONLY RE-PRIORITIZE and RE-PHRASE existing content to align with the JD's keywords and priorities. If a skill required by the JD is missing from the CV, do NOT add it; instead, highlight a transferable skill that IS in the CV.]
[END_MODIFIED_CV]

[START_TOP_FRAMEWORK]
[Provide ONLY the SINGLE highest-impact Strategic Initiative Framework here. Do NOT include ANY other sections or additional frameworks in this block.]
[END_TOP_FRAMEWORK]

[START_STRATEGY_PACKAGE]
[Include:]
1. Complete Gap Analysis (Parts 1.1–1.7 from the prompt).
2. TWO ADDITIONAL Strategic Frameworks (distinct from the Top Framework above).
3. TARGETED CONTACTS: Identify 3 decision-makers with LinkedIn Search URLs.
4. Follow-up timing guidance.
[END_STRATEGY_PACKAGE]

[START_OUTREACH_DOC]
[Outreach Email Templates for Word Document]
- Propose 2 outreach email templates addressed to specific roles identified in Targeted Contacts.
- CONTENT FOCUS: Briefly explain the "Top Strategic Initiative" plan you developed. Show how this specific plan addresses their pain points.
- AVOID: Empty boasting or generic "I am hard working" phrases. Focus on the VALUE of the plan.
[END_OUTREACH_DOC]

[START_EMAIL_BODY]
[This text will be the body of the email sent to the user Amanullah Khan]
[RECIPIENT ANALYSIS]: Identify the most likely primary recipient from your Targeted Contacts. Analyze their specific KPIs.

[ELITE EXECUTIVE OUTREACH]: Write a short, powerful 'Value Proposition' email for Amanullah Khan to use.
- ABSOLUTE LIMIT: 150 words.
- STYLE: Brief, punchy, and plan-centric.
- SUMMARY: Summarize the single most impactful insight from your Top Framework. Explain the 'How' (the plan) concisely.
- End with a low-friction question.
[END_EMAIL_BODY]

STYLE: 
- Human-written, elite strategist tone.
- NO markdown characters like # or * in your section dividers.
"""

    # ── Transparency: Save the prompt used ───────────────────────────────────
    (Path(__file__).parent / "prompts").mkdir(exist_ok=True)
    with open(Path(__file__).parent / "prompts" / "last_execution_prompt.txt", "w", encoding="utf-8") as f:
        f.write(full_prompt)

    # ── DeepSeek API call ────────────────────────────────────────────────────
    api_key = os.getenv("DEEPSEEK_API_KEY")
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")

    print("\n" + "=" * 50)
    print("  STEP 3 — DEEPSEEK AI PROCESSING (Detailed Analysis)")
    print("=" * 50)

    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "You are an elite career strategist. You provide polished, human-toned application strategies."},
                {"role": "user", "content": full_prompt},
            ],
            temperature=0.6,
            max_tokens=8000
        )
        ai_output = response.choices[0].message.content
    except Exception as e:
        print(f"\n❌ AI error: {e}")
        return

    # ── Save AI output for debugging ──────────────────────────────────────────
    with open(Path(__file__).parent / "prompts" / "last_ai_response.txt", "w", encoding="utf-8") as f:
        f.write(ai_output)

    # ── Split output with robust logic ────────────────────────────────────────
    def extract_section(text, start_tag, end_tag):
        # Handle variations: tags with/without brackets, optional slashes, and whitespace
        start_clean = start_tag.strip("[]")
        end_clean = end_tag.strip("[]")
        
        # Regex explanation:
        # (?:\[?\s*)? : non-capturing group for optional [ and whitespace
        # (?:\]?\s*)? : same for closing ]
        # (?i) : case-insensitive
        # (.*?) : non-greedy capture
        
        patterns = [
            # 1. Standard tags as requested: [TAG] CV [/TAG]
            fr"\[{re.escape(start_clean)}\](.*?)\[/?{re.escape(end_clean)}\]",
            # 2. Section headers: SECTION: CV ... SECTION: END_CV
            fr"{re.escape(start_clean)}[:\-\s]*(.*?)(?={re.escape(end_clean)}|$)",
            # 3. Just the tag name if it's on a line by itself
            fr"^{re.escape(start_clean)}[ \t]*\n(.*?)\n[/?]{re.escape(end_clean)}",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE | re.MULTILINE)
            if match and match.group(1).strip():
                content = match.group(1).strip()
                # Remove the AI's repetition of the instruction rule if it's there
                rule_text = "STRICT FACTUAL INTEGRITY RULE"
                if rule_text in content and len(content) < 500:
                    # If it's mostly the rule, this might be the wrong match or AI skipped it
                    continue
                return content
        return ""

    # Parse Metadata
    meta_block = extract_section(ai_output, "[METADATA]", "[/METADATA]")
    job_title = "Target Position"
    company = "Target Company"
    posted = "Unknown Date"
    
    if meta_block:
        for line in meta_block.splitlines():
            if "Title:" in line: job_title = line.split("Title:", 1)[1].strip()
            if "Company:" in line: company = line.split("Company:", 1)[1].strip()
            if "Posted:" in line: posted = line.split("Posted:", 1)[1].strip()

    cv_text_out = extract_section(ai_output, "[START_MODIFIED_CV]", "[END_MODIFIED_CV]")
    top_fw_text = extract_section(ai_output, "[START_TOP_FRAMEWORK]", "[END_TOP_FRAMEWORK]")
    package_text = extract_section(ai_output, "[START_STRATEGY_PACKAGE]", "[END_STRATEGY_PACKAGE]")
    ai_email_body = extract_section(ai_output, "[START_EMAIL_BODY]", "[END_EMAIL_BODY]")
    outreach_emails_doc = extract_section(ai_output, "[START_OUTREACH_DOC]", "[END_OUTREACH_DOC]")

    # ── Final Cleanup & Safety ────────────────────────────────────────────────
    # Scrub ALL possible tags from all extracted sections to prevent leakage
    all_sections = {
        "cv": cv_text_out,
        "top_fw": top_fw_text,
        "package": package_text,
        "email": ai_email_body,
        "outreach": outreach_emails_doc
    }
    
    tags_to_scrub = [
        "[METADATA]", "[/METADATA]",
        "[START_MODIFIED_CV]", "[END_MODIFIED_CV]", "[/END_MODIFIED_CV]", 
        "[START_TOP_FRAMEWORK]", "[END_TOP_FRAMEWORK]", "[/END_TOP_FRAMEWORK]",
        "[START_STRATEGY_PACKAGE]", "[END_STRATEGY_PACKAGE]", "[/END_STRATEGY_PACKAGE]",
        "[START_OUTREACH_DOC]", "[END_OUTREACH_DOC]", "[/END_OUTREACH_DOC]",
        "[START_EMAIL_BODY]", "[END_EMAIL_BODY]", "[/END_EMAIL_BODY]"
    ]

    for key in all_sections:
        if all_sections[key]:
            for tag in tags_to_scrub:
                all_sections[key] = all_sections[key].replace(tag, "").strip()
            # If a START tag from another section leaked in, truncate there
            if "[START_" in all_sections[key]:
                all_sections[key] = all_sections[key].split("[START_")[0].strip()

    # Re-assign cleaned values
    cv_text_out = all_sections["cv"]
    top_fw_text = all_sections["top_fw"]
    package_text = all_sections["package"]
    ai_email_body = all_sections["email"]
    outreach_emails_doc = all_sections["outreach"]

    if not cv_text_out or not top_fw_text:
        print("\n  ⚠️  STRUCTURE ERROR: Primary sections were not detected.")
        print(f"     CV Detected: {'Yes' if cv_text_out else 'NO'}")
        print(f"     Top Strategy Detected: {'Yes' if top_fw_text else 'NO'}")
        print("     Attempting heuristic recovery from raw output...")
        
        # Split by ANY tag-like bracketed uppercase word
        parts = re.split(r'\[[A-Z_ /]+\]', ai_output)
        valid_parts = [p.strip() for p in parts if len(p.strip()) > 150]
        
        if len(valid_parts) >= 3:
            cv_text_out = cv_text_out or valid_parts[0]
            top_fw_text = top_fw_text or valid_parts[1]
            package_text = package_text or valid_parts[2]
            if not ai_email_body and len(valid_parts) >= 4:
                ai_email_body = valid_parts[3]

    # ── Build Word docs ──────────────────────────────────────────────────────
    print("\n  Building formatted Word documents…")
    company_slug = company.replace(' ', '_')
    
    f1 = build_docx("",  cv_text_out,    f"Modified_CV_{company_slug}.docx")
    f2 = build_docx("Top Strategic Initiative",  top_fw_text,    f"Top_Strategy_{company_slug}.docx")
    f3 = build_docx("Full Career Strategy Pack", package_text,   f"Full_Strategy_Package_{company_slug}.docx")
    
    attachments = [f1, f2, f3]
    
    if outreach_emails_doc:
        f4 = build_docx("Outreach Email Templates", outreach_emails_doc, f"Outreach_Emails_{company_slug}.docx")
        attachments.append(f4)
        print(f"  ✅ Saved: {Path(f4).name}")
    
    print(f"  ✅ Saved: {Path(f1).name}")
    print(f"  ✅ Saved: {Path(f2).name}")
    print(f"  ✅ Saved: {Path(f3).name}")

    # ── Email ────────────────────────────────────────────────────────────────
    print("\n" + "=" * 50)
    print("  STEP 4 — SEND VIA EMAIL")
    print("=" * 50)
    print("  Press Enter to use the default: aa.khan.9093@gmail.com")
    recipient = input("  Recipient email: ").strip() or "aa.khan.9093@gmail.com"

    # BUILD SUBJECT: job title, company and the date of job posting
    subject = f"Package: {job_title} | {company} | Posted: {posted}"
    
    print(f"  Sending to {recipient}…")
    ok = send_job_bot_email(recipient, subject, ai_email_body, attachments)
    if ok:
        print("\n  ✅ Done! Files sent to your inbox.")
    else:
        print("\n  ❌ Email failed. Files are saved locally in the project folder.")


if __name__ == "__main__":
    run()
