import streamlit as st
import os
import re
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup
import datetime
from pypdf import PdfReader
import io

# Load environment logic
load_dotenv()

st.set_page_config(page_title="Strategic Career Transition Bot", page_icon="🚀", layout="wide")

# ────────────────────────────────────────────────────────────────────────────
# HELPERS
# ────────────────────────────────────────────────────────────────────────────

def scrape_url(url: str) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n").strip()
    except Exception as e:
        return ""

def process_line_with_formatting(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
        else:
            clean_part = part.replace("*", "").strip()
            if clean_part:
                paragraph.add_run(f" {clean_part} " if part != parts[0] else clean_part)

def build_docx_in_memory(title: str, content: str) -> io.BytesIO:
    doc = Document()
    DEEP_BLUE = RGBColor(0, 51, 102)
    DARK_GREY = RGBColor(64, 64, 64)
    style = doc.styles["Normal"]
    style.font.name = "Outfit" if "Outfit" in [f.name for f in doc.styles] else "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GREY

    if title:
        h0 = doc.add_heading(title.upper(), level=0)
        h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h0.runs:
            run.font.color.rgb = DEEP_BLUE
        doc.add_paragraph()

    content = content.replace("---FILE_SPLIT---", "").strip()
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].replace("#", "").strip(), level=3)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].replace("#", "").strip(), level=2)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].replace("#", "").strip(), level=1)
            for run in p.runs: run.font.color.rgb = DEEP_BLUE
        elif line.startswith(("- ", "• ", "* ")) and not line.startswith("**"):
            p = doc.add_paragraph(style="List Bullet")
            process_line_with_formatting(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            process_line_with_formatting(p, line)
            
    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

def extract_section(text, start_tag, end_tag):
    start_clean = start_tag.strip("[]")
    end_clean = end_tag.strip("[]")
    patterns = [
        fr"\[{re.escape(start_clean)}\](.*?)\[/?{re.escape(end_clean)}\]",
        fr"{re.escape(start_clean)}[:\-\s]*(.*?)(?={re.escape(end_clean)}|$)",
        fr"^{re.escape(start_clean)}[ \t]*\n(.*?)\n[/?]{re.escape(end_clean)}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE | re.MULTILINE)
        if match and match.group(1).strip():
            content = match.group(1).strip()
            if "STRICT FACTUAL INTEGRITY RULE" in content and len(content) < 500:
                continue
            return content
    return ""

# ────────────────────────────────────────────────────────────────────────────
# UI
# ────────────────────────────────────────────────────────────────────────────

st.title("🚀 Strategic Career Transition Bot")
st.markdown("Automate your targeted job application materials using elite career strategy frameworks.")

st.markdown("---")

col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Job Description")
    jd_url = st.text_input("Job Description URL (We will try to scrape it)")
    jd_paste = st.text_area("Or paste the Job Description text directly", height=200)

with col2:
    st.subheader("2. Your CV")
    cv_file = st.file_uploader("Upload your CV (PDF or TXT/MD)", type=["pdf", "txt", "md"])
    cv_paste = st.text_area("Or paste your CV text directly", height=200)

st.markdown("---")

if st.button("Generate Strategy Package", type="primary", use_container_width=True):
    # Gather Data
    jd_text = ""
    if jd_url:
        with st.spinner("Scraping Job Description..."):
            jd_text = scrape_url(jd_url)
    if not jd_text and jd_paste:
        jd_text = jd_paste.strip()
        
    cv_text = ""
    if cv_file:
        if cv_file.name.endswith(".pdf"):
            reader = PdfReader(cv_file)
            for page in reader.pages:
                cv_text += page.extract_text() + "\n"
        else:
            cv_text = cv_file.getvalue().decode("utf-8")
    if not cv_text and cv_paste:
        cv_text = cv_paste.strip()
        
    if not jd_text or not cv_text:
        st.error("Please provide both a Job Description and your CV.")
    else:
        with st.spinner("Analyzing profile and generating strategy package... This usually takes 1-2 minutes."):
            try:
                # Load prompt logic
                today = datetime.date.today().strftime("%B %d, %Y")
                prompt_file = Path(__file__).parent / "prompts" / "career_transition_logic.md"
                with open(prompt_file, "r", encoding="utf-8") as f:
                    prompt_logic = f.read()

                full_prompt = f"""{prompt_logic}

===== CONTEXT =====
TODAY'S DATE: {today}

===== INPUT 1 — JOB DESCRIPTION =====
{jd_text}

===== INPUT 2 — CANDIDATE CV =====
{cv_text}

===== OUTPUT FORMAT INSTRUCTIONS (CRITICAL) =====
You MUST separate your response into FIVE distinct sections using these EXACT labels. Do not nest them. Do not include extra text outside these tags.

[METADATA]
Title: [Extract Job Title]
Company: [Extract Company Name]
Posted: [Look for the date of posting in the JD. If not found, reason from content or use '{today}']
[/METADATA]

[START_MODIFIED_CV]
[STRICT FACTUAL INTEGRITY RULE: Provide ONLY the fully tailored Modified CV here. You MUST NOT invent, make up, or hallucinate any skills, companies, or experiences that are not present in Input 2 (Candidate CV). Your task is to ONLY RE-PRIORITIZE and RE-PHRASE existing content to align with the JD's keywords and priorities. If a skill required by the JD is missing from the CV, do NOT add it; instead, highlight a transferable skill that IS in the CV.]
[END_MODIFIED_CV]

[START_TOP_FRAMEWORK]
[Provide ONLY the SINGLE highest-impact Strategic Initiative Framework here. Do NOT include ANY other sections or additional frameworks in this block.]
[END_TOP_FRAMEWORK]

[START_STRATEGY_PACKAGE]
[Include:]
1. Complete Gap Analysis (Parts 1.1–1.7 from the prompt).
2. TWO ADDITIONAL Strategic Frameworks (distinct from the Top Framework above).
3. TARGETED CONTACTS: Identify 3 decision-makers with LinkedIn Search URLs.
4. Follow-up timing guidance.
[END_STRATEGY_PACKAGE]

[START_OUTREACH_DOC]
[Outreach Email Templates for Word Document]
- Propose 2 outreach email templates addressed to specific roles identified in Targeted Contacts.
- CONTENT FOCUS: Briefly explain the "Top Strategic Initiative" plan you developed. Show how this specific plan addresses their pain points.
- AVOID: Empty boasting or generic "I am hard working" phrases. Focus on the VALUE of the plan.
[END_OUTREACH_DOC]

[START_EMAIL_BODY]
[This text will be the body of the email sent to the user Amanullah Khan]
[RECIPIENT ANALYSIS]: Identify the most likely primary recipient from your Targeted Contacts. Analyze their specific KPIs.

[ELITE EXECUTIVE OUTREACH]: Write a short, powerful 'Value Proposition' email for Amanullah Khan to use.
- ABSOLUTE LIMIT: 150 words.
- STYLE: Brief, punchy, and plan-centric.
- SUMMARY: Summarize the single most impactful insight from your Top Framework. Explain the 'How' (the plan) concisely.
- End with a low-friction question.
[END_EMAIL_BODY]

STYLE: 
- Human-written, elite strategist tone.
- NO markdown characters like # or * in your section dividers.
"""
                api_key = os.getenv("DEEPSEEK_API_KEY")
                if not api_key:
                    st.error("Missing DEEPSEEK_API_KEY in environment or .env file.")
                    st.stop()
                    
                client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": "You are an elite career strategist. You provide polished, human-toned application strategies."},
                        {"role": "user", "content": full_prompt},
                    ],
                    temperature=0.6,
                    max_tokens=8000
                )
                ai_output = response.choices[0].message.content
                
                # Parse output
                meta_block = extract_section(ai_output, "[METADATA]", "[/METADATA]")
                company = "Target_Company"
                if meta_block:
                    for line in meta_block.splitlines():
                        if "Company:" in line: company = line.split("Company:", 1)[1].strip()

                cv_text_out = extract_section(ai_output, "[START_MODIFIED_CV]", "[END_MODIFIED_CV]")
                top_fw_text = extract_section(ai_output, "[START_TOP_FRAMEWORK]", "[END_TOP_FRAMEWORK]")
                package_text = extract_section(ai_output, "[START_STRATEGY_PACKAGE]", "[END_STRATEGY_PACKAGE]")
                ai_email_body = extract_section(ai_output, "[START_EMAIL_BODY]", "[END_EMAIL_BODY]")
                outreach_emails_doc = extract_section(ai_output, "[START_OUTREACH_DOC]", "[END_OUTREACH_DOC]")
                
                # Scrubing
                all_sections = {
                    "cv": cv_text_out, "top_fw": top_fw_text, "package": package_text,
                    "email": ai_email_body, "outreach": outreach_emails_doc
                }
                tags_to_scrub = [
                    "[METADATA]", "[/METADATA]",
                    "[START_MODIFIED_CV]", "[END_MODIFIED_CV]", "[/END_MODIFIED_CV]", 
                    "[START_TOP_FRAMEWORK]", "[END_TOP_FRAMEWORK]", "[/END_TOP_FRAMEWORK]",
                    "[START_STRATEGY_PACKAGE]", "[END_STRATEGY_PACKAGE]", "[/END_STRATEGY_PACKAGE]",
                    "[START_OUTREACH_DOC]", "[END_OUTREACH_DOC]", "[/END_OUTREACH_DOC]",
                    "[START_EMAIL_BODY]", "[END_EMAIL_BODY]", "[/END_EMAIL_BODY]"
                ]
                for key in all_sections:
                    if all_sections[key]:
                        for tag in tags_to_scrub:
                            all_sections[key] = all_sections[key].replace(tag, "").strip()
                        if "[START_" in all_sections[key]:
                            all_sections[key] = all_sections[key].split("[START_")[0].strip()

                cv_text_out = all_sections["cv"]
                top_fw_text = all_sections["top_fw"]
                package_text = all_sections["package"]
                ai_email_body = all_sections["email"]
                outreach_emails_doc = all_sections["outreach"]

                if not cv_text_out or not top_fw_text:
                    parts = re.split(r'\[[A-Z_ /]+\]', ai_output)
                    valid_parts = [p.strip() for p in parts if len(p.strip()) > 150]
                    if len(valid_parts) >= 3:
                        cv_text_out = cv_text_out or valid_parts[0]
                        top_fw_text = top_fw_text or valid_parts[1]
                        package_text = package_text or valid_parts[2]
                        if not ai_email_body and len(valid_parts) >= 4:
                            ai_email_body = valid_parts[3]

                st.success("Analysis Complete! Your documents are ready for download.")
                
                company_slug = company.replace(' ', '_')
                
                st.subheader("📥 Download Your Materials")
                dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
                
                if cv_text_out:
                    doc1 = build_docx_in_memory("", cv_text_out)
                    dl_col1.download_button("📄 Modified CV", data=doc1, file_name=f"Modified_CV_{company_slug}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                if top_fw_text:
                    doc2 = build_docx_in_memory("Top Strategic Initiative", top_fw_text)
                    dl_col2.download_button("🚀 Top Strategy", data=doc2, file_name=f"Top_Strategy_{company_slug}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                if package_text:
                    doc3 = build_docx_in_memory("Full Career Strategy Pack", package_text)
                    dl_col3.download_button("📊 Full Strategy Pack", data=doc3, file_name=f"Full_Strategy_Package_{company_slug}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
                if outreach_emails_doc:
                    doc4 = build_docx_in_memory("Outreach Email Templates", outreach_emails_doc)
                    dl_col4.download_button("✉️ Outreach Emails", data=doc4, file_name=f"Outreach_Emails_{company_slug}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                st.markdown("---")
                st.subheader("Quick Review")
                with st.expander("Top Strategic Initiative Preview", expanded=True):
                    st.markdown(top_fw_text)
                with st.expander("Modified CV Preview", expanded=False):
                    st.markdown(cv_text_out)
                with st.expander("Email Body Proposal", expanded=False):
                    st.markdown(ai_email_body)

            except Exception as e:
                st.error(f"Error during API call or processing: {str(e)}")
